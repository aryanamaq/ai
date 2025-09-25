from docx import Document
from pathlib import Path
from typing import List, Dict, Any, Optional
import json
import hashlib
from datetime import datetime
import pandas as pd
from docx.document import Document as DocxDocument
from docx.table import Table
import re
from schemas import DocumentType, ValidationStatus, DocumentExtraction

class DocumentProcessor:
    def __init__(self, upload_dir: Path):
        self.upload_dir = upload_dir
        self.upload_dir.mkdir(exist_ok=True)

    def extract_text_from_word(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from Word document."""
        try:
            doc: DocxDocument = Document(file_path)
            
            # Extract text with section markers
            sections = self._extract_sections(doc)
            
            # Extract tables
            tables = self._extract_tables(doc)
            
            # Extract signatures
            signatures = self._detect_signatures(doc)
            
            # Identify document type
            doc_type = self._classify_document(sections)
            
            return {
                "sections": sections,
                "tables": tables,
                "signatures": signatures,
                "document_type": doc_type
            }
        except Exception as e:
            raise ValueError(f"Error processing Word document: {str(e)}")

    def _extract_sections(self, doc: DocxDocument) -> Dict[str, str]:
        """Extract document sections based on headings and structure."""
        sections = {}
        current_section = "header"
        current_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a heading
            if para.style.name.startswith('Heading'):
                if current_text:
                    sections[current_section] = '\n'.join(current_text)
                current_section = text.lower().replace(' ', '_')
                current_text = []
            else:
                current_text.append(text)

        # Add the last section
        if current_text:
            sections[current_section] = '\n'.join(current_text)

        return sections

    def _extract_tables(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract and structure tables from the document."""
        tables = []
        for table in doc.tables:
            try:
                # Convert table to DataFrame
                data = [[cell.text for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data[1:], columns=data[0])
                
                # Convert DataFrame to dict
                table_dict = {
                    "headers": data[0],
                    "data": df.to_dict('records'),
                    "location": f"Table {len(tables) + 1}"
                }
                tables.append(table_dict)
            except Exception:
                continue
        return tables

    def _detect_signatures(self, doc: DocxDocument) -> Dict[str, bool]:
        """Detect signature blocks and validate their presence."""
        signature_indicators = [
            r"signature",
            r"signed by",
            r"authorized signatory",
            r"in witness whereof"
        ]
        
        text = "\n".join([p.text.lower() for p in doc.paragraphs])
        
        signatures = {
            "has_signature_block": any(re.search(pattern, text) for pattern in signature_indicators),
            "signature_locations": []
        }
        
        # Find signature locations
        for i, para in enumerate(doc.paragraphs):
            if any(re.search(pattern, para.text.lower()) for pattern in signature_indicators):
                signatures["signature_locations"].append(f"Paragraph {i + 1}")
        
        return signatures

    def _classify_document(self, sections: Dict[str, str]) -> DocumentType:
        """Classify document type based on content analysis."""
        text = " ".join(sections.values()).lower()
        
        # Define classification rules
        rules = {
            DocumentType.LEASE: ["lease", "tenant", "landlord", "rent", "property"],
            DocumentType.LOAN: ["loan", "borrower", "lender", "interest rate", "principal"],
            DocumentType.MSA: ["services", "statement of work", "sla", "deliverables"]
        }
        
        # Score each document type
        scores = {doc_type: sum(1 for keyword in keywords if keyword in text)
                 for doc_type, keywords in rules.items()}
        
        # Return the type with highest score
        return max(scores.items(), key=lambda x: x[1])[0]

    def validate_document(self, doc_data: Dict[str, Any]) -> Dict[str, ValidationStatus]:
        """Validate document contents against business rules."""
        validations = {}
        
        # Required sections
        required_sections = ["header", "parties", "terms"]
        validations["required_sections"] = ValidationStatus.PASSED if all(
            section in doc_data["sections"] for section in required_sections
        ) else ValidationStatus.FAILED
        
        # Signature validation
        validations["signatures"] = ValidationStatus.PASSED if doc_data["signatures"]["has_signature_block"] else ValidationStatus.WARNING
        
        # Date order validation
        if "effective_date" in doc_data and "execution_date" in doc_data:
            if doc_data["effective_date"] < doc_data["execution_date"]:
                validations["date_order"] = ValidationStatus.WARNING
            else:
                validations["date_order"] = ValidationStatus.PASSED
        
        return validations

    def save_document(self, file_content: bytes, filename: str) -> Dict[str, str]:
        """Save document with metadata and validation."""
        try:
            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_filename = f"{timestamp}_{self._sanitize_filename(filename)}"
            file_path = self.upload_dir / safe_filename
            
            # Save file
            with open(file_path, "wb") as f:
                f.write(file_content)
            
            # Generate metadata
            metadata = {
                "document_id": self._generate_document_id(file_content),
                "original_filename": filename,
                "saved_filename": safe_filename,
                "upload_time": datetime.now().isoformat(),
                "file_size": len(file_content),
                "file_hash": self._calculate_hash(file_content),
                "processing_status": "pending"
            }
            
            # Save metadata
            meta_path = file_path.with_suffix('.meta.json')
            with open(meta_path, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            return metadata
        except Exception as e:
            raise ValueError(f"Error saving document: {str(e)}")

    def get_document_list(self) -> List[Dict[str, Any]]:
        """Get list of documents with metadata and processing status."""
        documents = []
        for meta_file in self.upload_dir.glob("*.meta.json"):
            try:
                with open(meta_file, 'r') as f:
                    metadata = json.load(f)
                doc_path = meta_file.with_suffix('').with_suffix('.docx')
                analysis_path = meta_file.with_suffix('.analysis.json')
                
                if doc_path.exists():
                    metadata["exists"] = True
                    metadata["path"] = str(doc_path)
                    metadata["has_analysis"] = analysis_path.exists()
                    if metadata["has_analysis"]:
                        with open(analysis_path, 'r') as f:
                            analysis = json.load(f)
                            metadata["validation_summary"] = analysis.get("validation_results", {})
                    documents.append(metadata)
            except Exception:
                continue
        return sorted(documents, key=lambda x: x["upload_time"], reverse=True)

    def export_to_csv(self, document_id: str) -> str:
        """Export document analysis to CSV."""
        try:
            analysis_path = next(self.upload_dir.glob(f"*{document_id}*.analysis.json"))
            with open(analysis_path, 'r') as f:
                analysis = json.load(f)
            
            # Convert nested JSON to flat structure
            flat_data = self._flatten_json(analysis)
            df = pd.DataFrame([flat_data])
            
            # Save to CSV
            csv_path = analysis_path.with_suffix('.csv')
            df.to_csv(csv_path, index=False)
            return str(csv_path)
        except Exception as e:
            raise ValueError(f"Error exporting to CSV: {str(e)}")

    def _flatten_json(self, nested_json: Dict, parent_key: str = '', sep: str = '_') -> Dict:
        """Flatten nested JSON structure."""
        items = []
        for k, v in nested_json.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_json(v, new_key, sep).items())
            else:
                items.append((new_key, v))
        return dict(items)

    def _generate_document_id(self, content: bytes) -> str:
        """Generate unique document ID."""
        return hashlib.sha256(content).hexdigest()[:12]

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe storage."""
        return "".join(c for c in filename if c.isalnum() or c in "._- ")

    def _calculate_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of file content."""
        return hashlib.sha256(content).hexdigest()

    def validate_file(self, filename: str, content: bytes) -> None:
        """Validate file before processing."""
        if not filename.lower().endswith(('.doc', '.docx')):
            raise ValueError("Only Word documents (.doc, .docx) are supported")
        
        if len(content) > 10 * 1024 * 1024:  # 10MB limit
            raise ValueError("File size exceeds maximum limit of 10MB")
